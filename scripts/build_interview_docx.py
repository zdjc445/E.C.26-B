from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "识价镜项目八股考点手册.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x55, 0x65, 0x74)
GRID = "C9D3DF"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "识价镜项目八股考点手册"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.runs[0], 9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer_p.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r2 = footer_p.add_run()
    r2._r.append(fld_begin)
    r2._r.append(instr)
    r2._r.append(fld_end)
    for run in footer_p.runs:
        set_run_font(run, 9, color=MUTED)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("项目面试复习手册")
    set_run_font(r, 11, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("识价镜项目八股考点手册")
    set_run_font(r, 25, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("覆盖 Spring Boot 后端、Flutter Android、聊天式购物 Agent、Ark/Mock AI 编排、检索推荐、数据与 Mock、测试验收、答辩追问和已有八股迁移。")
    set_run_font(r, 11.5, color=MUTED)

    rows = [
        ("项目", "识价镜"),
        ("定位", "面向 C 端消费者的聊天式购物 Agent"),
        ("当前阶段", "聊天式 AI 识别 + 公开样例数据多平台比价 + 7 维度自然语言筛选 + 动态建议卡 + 持久化 + 认证 + 收藏 + 价格提醒 + 语音转写阶段"),
        ("生成日期", "2026-07-08"),
        ("资料来源", "本地 README、docs/00-14、backend 源码与测试、app 源码与测试、八股资料目录"),
    ]
    add_kv_table(doc, rows, widths=(1.35, 5.15))

    add_callout(doc, "使用方法",
                "先背“项目一句话”和“核心创新点”，再按章节准备追问。每个考点都绑定项目落点：面试时先讲通用原理，再指出本项目中对应的类、接口、字段或配置。")

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 11)
    else:
        r = p.add_run(text)
        set_run_font(r, 11)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 11)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, 11)


def set_cell(cell, text: str, bold: bool = False, fill: str | None = None,
             font_size: float = 9.5, color: RGBColor = INK) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.text = ""
    r = p.add_run(text)
    set_run_font(r, font_size, bold=bold, color=color)
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(i, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), str(val))
                el.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, fill=HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], str(value))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def add_kv_table(doc: Document, rows: Sequence[tuple[str, str]], widths: tuple[float, float]) -> None:
    add_table(
        doc,
        ("项", "内容"),
        rows,
        (int(widths[0] * 1440), int(widths[1] * 1440)),
    )


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, (CONTENT_WIDTH_DXA,), indent_dxa=TABLE_INDENT_DXA)
    cell = table.cell(0, 0)
    set_cell(cell, f"{title}：{body}", bold=False, fill=CALLOUT_FILL, font_size=10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def add_qa(doc: Document, q: str, answer: str, project: str, follow: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("问：")
    set_run_font(r, 10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(q)
    set_run_font(r, 10.5, bold=True, color=INK)
    add_para(doc, "答：" + answer, bold_prefix="答：")
    add_para(doc, "项目落点：" + project, bold_prefix="项目落点：")
    if follow:
        add_para(doc, "追问：" + follow, bold_prefix="追问：")


def build_doc() -> None:
    doc = Document()
    style_doc(doc)
    add_footer(doc)
    add_title_page(doc)

    add_heading(doc, "0. 先背这页", 1)
    add_callout(
        doc,
        "项目一句话",
        "识价镜是一个 Flutter Android + Spring Boot Java 21 的聊天式购物 Agent，用 Mock 保证闭环，用 Ark 做图片识别、意图解析和解释改写增强，完成拍照识物、自然语言多轮筛选、公开样例商品检索、多平台样例报价、同款分组、收藏、价格提醒和语音转写。"
    )
    add_table(doc, ("面试表达", "必须说清楚的点"), [
        ("核心创新", "三段式 AI 编排（规则默认 + Ark 增强 + 失败回退）、taxonomy 检索归一、跨轮 7 维自然语言筛选、动态建议卡。"),
        ("工程边界", "商品搜索默认读取公开样例数据，不调用真实电商接口；平台报价和价格历史由系统生成，不代表真实平台。"),
        ("AI 边界", "AI 输出结构化识别/解析/解释摘要；商品 ID、平台、价格、排序、数值分数由规则和数据链路控制。"),
        ("可验证性", "文档记录后端 142 tests 通过；Flutter test 41 tests 通过；测试覆盖多轮筛选、Ark 回退、识别修正、历史恢复、收藏和价格提醒。"),
    ], (2200, 7160))

    add_heading(doc, "1. 项目全景与技术栈", 1)
    add_heading(doc, "1.1 目录和模块", 2)
    add_table(doc, ("目录", "职责", "面试价值"), [
        ("backend/", "Spring Boot API 服务；包名 com.ec26b.shoppingagent。", "后端工程、REST、认证、持久化、AI Provider、推荐检索。"),
        ("app/", "Flutter Android 客户端；聊天首页、历史抽屉、我的页、收藏、价格提醒、偏好记忆。", "Flutter 状态管理、路由、网络、多媒体、移动端本地存储。"),
        ("docs/", "项目范围、架构、API、Agent、数据、测试、答辩、AI 使用总结。", "答辩口径和业务边界来源。"),
        ("scripts/", "build_public_product_offers.py 等辅助脚本。", "数据确定性构建、可复现实验。"),
        ("outputs/", "答辩 PPT、演示素材和本手册输出。", "交付物管理。"),
        ("八股/", "已有 Java/MySQL/Redis/Spring 等八股资料和前端浏览器页面。", "通用八股补充库。"),
    ], (1800, 3900, 3660))

    add_heading(doc, "1.2 精确技术栈", 2)
    add_table(doc, ("层级", "技术/版本", "项目落点"), [
        ("客户端", "Flutter SDK >=3.22.0；Dart SDK ^3.4.0", "app/pubspec.yaml。"),
        ("客户端 resolved 依赖", "flutter_riverpod 2.6.1；go_router 14.8.1；http 1.6.0；image_picker 1.2.1；shared_preferences 2.5.3", "app/pubspec.lock。"),
        ("Android 构建", "Android Gradle Plugin 8.7.3；Kotlin 2.1.0；Gradle wrapper 8.12；ndkVersion 27.0.12077973", "app/android/settings.gradle.kts、gradle-wrapper.properties、app/build.gradle.kts。"),
        ("Android 包名", "namespace/applicationId = com.ec26b.shopping_agent_app", "app/android/app/build.gradle.kts。"),
        ("客户端状态", "flutter_riverpod ^2.5.0", "apiBaseUrlProvider、chatControllerProvider、userProfileProvider。"),
        ("客户端路由", "go_router ^14.0.0", "路由 /login、/home、/me、/favorites、/price-alerts、/preferences。"),
        ("客户端网络", "http ^1.2.0", "REST JSON 与 Multipart 上传。"),
        ("客户端图片", "image_picker ^1.1.0", "拍照/相册入口。"),
        ("客户端本地存储", "shared_preferences ^2.3.0", "MemoryStore 保存 profile/events/flags。"),
        ("后端语言", "Java 21", "pom.xml 的 java.version=21。"),
        ("后端框架", "Spring Boot 3.4.1", "spring-boot-starter-web/validation/jdbc/security/test。"),
        ("后端受管依赖解析版本", "Spring Framework 6.2.1；PostgreSQL 42.7.4；Flyway 10.20.1；spring-security-test 6.4.2；JUnit Jupiter 5.11.4", "本机 Maven 缓存和 Spring Boot 3.4.1 dependency management。"),
        ("构建", "Maven；spring-boot-maven-plugin；maven-compiler-plugin release=21", "backend/pom.xml。"),
        ("认证", "Spring Security + JJWT 0.12.6 + BCryptPasswordEncoder", "SecurityConfig、JwtService、PasswordHasher。"),
        ("数据库", "PostgreSQL runtime + Flyway", "postgres profile 开启 datasource 与 db/migration/V1__init.sql。"),
        ("AI Provider", "Mock / Ark；base URL 默认 https://ark.cn-beijing.volces.com/api/v3", "AiConfig、ArkClient、VoiceConfig。"),
    ], (1700, 3300, 4360))

    add_heading(doc, "1.3 配置键和环境变量", 2)
    add_table(doc, ("配置键/环境变量", "默认值", "作用"), [
        ("server.port", "8080", "后端服务端口。"),
        ("app.upload-dir / UPLOAD_DIR", "../uploads", "图片上传落盘目录。"),
        ("app.persistence.store / APP_PERSISTENCE_STORE", "memory", "memory / postgres 仓储切换。"),
        ("app.product-source.mode / PRODUCT_SOURCE_MODE", "public-dataset-platforms", "商品源模式。"),
        ("app.product-source.public-resource / PUBLIC_PRODUCT_RESOURCE", "data/public-product-offers.json", "公开样例商品资源。"),
        ("app.ai.provider / AI_PROVIDER", "mock", "mock / ark AI Provider 切换。"),
        ("app.ai.ark.api-key / ARK_API_KEY", "", "Ark API Key。"),
        ("app.ai.ark.endpoint-id / ARK_ENDPOINT_ID", "", "Ark Endpoint ID，作为 chat/completions 的 model。"),
        ("app.ai.ark.base-url / ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3", "Ark 基础 URL。"),
        ("app.auth.enabled / AUTH_ENABLED", "false", "是否强制要求登录。"),
        ("app.auth.jwt-secret / AUTH_JWT_SECRET", "dev-only-256bit-secret-please-change-me-12345678", "JWT HMAC 密钥。"),
        ("app.auth.jwt-ttl-minutes / AUTH_JWT_TTL_MINUTES", "1440", "JWT 过期分钟数。"),
        ("app.voice.provider / VOICE_PROVIDER", "mock", "语音 Provider。"),
        ("app.voice.transcript-fallback / VOICE_FALLBACK_TEXT", "推荐运动鞋", "Mock 语音转写文本。"),
        ("EC26B_API_BASE_URL", "http://localhost:8080", "Flutter 编译期后端地址。"),
    ], (3200, 2600, 3760))

    add_heading(doc, "2. 后端八股考点", 1)
    add_heading(doc, "2.1 Spring Boot 分层和包职责", 2)
    add_table(doc, ("包", "核心类", "需要会讲的八股"), [
        ("api", "HealthController、ImageController、ChatController、RecognitionController、AuthController、FavoriteController、PriceAlertController、VoiceController、EcommerceController", "REST 设计、统一响应、参数校验、错误码、multipart、Controller 边界。"),
        ("ai", "AiRecognitionProvider、MockRecognitionProvider、ArkRecognitionProvider、FallbackRecognitionProvider、ArkClient、RecognitionStore", "策略模式、回退、外部服务调用、结构化输出、内存存储。"),
        ("chat", "ChatHistoryRepository、InMemoryChatHistoryRepository、PostgresChatHistoryRepository、ChatStore、MockAgent", "会话模型、历史恢复、内存/数据库仓储切换、Agent 编排。"),
        ("product", "CategoryResolver、CompositeProductSourceProvider、PublicDatasetProductSourceProvider、RecommendationScorer、RecommendationExplainer、ArkShoppingIntentParser、FallbackShoppingIntentParser", "检索、排序、同款分组、可解释推荐、Prompt、RAG 扩展。"),
        ("auth", "AuthService、CurrentUser、JwtService、PasswordHasher、SecurityConfig、UserRepository", "JWT、BCrypt、无状态认证、Controller 内权限边界。"),
        ("favorite/alert", "FavoriteRepository、PriceAlertRepository 及 memory/postgres 实现", "仓储接口、条件装配、用户维度数据隔离。"),
        ("voice", "VoiceTranscriber、MockVoiceTranscriber、ArkVoiceTranscriber、FallbackVoiceTranscriber", "语音接口、Provider 切换、回退。"),
    ], (1500, 3600, 4260))
    add_callout(doc, "启动补充", "ShoppingAgentApplication 启动 Spring Boot，并读取 .env / ../.env 到 System.setProperty；agent、persistence、security、service 目录当前只有 .gitkeep。")

    add_heading(doc, "2.2 REST API 清单", 2)
    add_table(doc, ("方法", "路径", "用途"), [
        ("GET", "/api/health", "健康状态、stage、aiProvider、persistenceStore、authEnabled、ecommerceProvider、voiceProvider。"),
        ("POST", "/api/auth/register", "注册；请求 username、password、displayName。"),
        ("POST", "/api/auth/login", "登录；请求 username、password。"),
        ("GET", "/api/auth/me", "当前用户；Authorization: Bearer token 可选，AUTH_ENABLED=false 时 demo 用户兜底。"),
        ("POST", "/api/images/upload", "multipart/form-data，文件字段名 file。"),
        ("POST", "/api/chat/sessions", "创建聊天会话。"),
        ("GET", "/api/chat/sessions", "会话摘要列表。"),
        ("POST", "/api/chat/sessions/{sessionId}/messages", "发送文本/图片/选项，返回 AgentReply。"),
        ("GET", "/api/chat/sessions/{sessionId}/messages", "历史消息，assistant 消息含 agentReply。"),
        ("PATCH", "/api/chat/sessions/{sessionId}", "重命名会话，请求 title。"),
        ("DELETE", "/api/chat/sessions/{sessionId}", "删除会话。"),
        ("POST", "/api/recognition", "图片识别，请求 imageId。"),
        ("PATCH", "/api/recognition/{recognitionId}/attributes", "修正 category、brand、model、attributes。"),
        ("GET", "/api/ecommerce/status", "商品源状态和样例平台/品类。"),
        ("POST", "/api/favorites", "新增收藏。"),
        ("GET", "/api/favorites", "收藏列表。"),
        ("DELETE", "/api/favorites/{productId}", "删除收藏。"),
        ("POST", "/api/price-alerts", "创建价格提醒。"),
        ("GET", "/api/price-alerts", "价格提醒列表。"),
        ("DELETE", "/api/price-alerts/{alertId}", "删除价格提醒。"),
        ("POST", "/api/price-alerts/check", "用当前 Mock 价格检测提醒。"),
        ("POST", "/api/voice/transcribe", "multipart/form-data，文件字段名 file。"),
    ], (900, 3550, 4910))

    add_heading(doc, "2.3 统一响应和关键 DTO", 2)
    add_table(doc, ("对象", "字段"), [
        ("ApiResponse<T>", "code、message、data。成功 code=0、message=success。"),
        ("AuthResult", "token、userId、username、displayName、role、expiresInSeconds。"),
        ("ChatMessageRequest", "text、imageIds、selectedOptionIds、profile。"),
        ("ImagePayload", "imageId、contentType、bytes、filename。"),
        ("ImageMetadata", "imageId、storedFileName、originalFileName、contentType、size、createdAt。"),
        ("RecognitionResult", "imageId、category、brand、model、keywords、attributes、confidence、aiProvider、fallbackUsed、explanation、notices。"),
        ("ProductSearchQuery", "keyword、preferences、maxPrice、color、brand、platforms、sortBy、minRating、profile。"),
        ("ProductOffer", "productId、title、platform、price、originalPrice、shopName、imageUrl、productUrl、rating、sales、tags、reasons、score、brand、priceHistory、matchedPreferences、sameItemKey。"),
        ("AgentReply", "replyId、replyType、text、cards。"),
        ("Card", "cardType、title、products、platformStats、decisionScore、decisionSignals、evidence、risks、productAnalyses、intentProvider、intentFallbackUsed、explanationProvider、explanationFallbackUsed、notices、filterSummary、groups、emptyReason 等。"),
        ("ProductGroup", "groupId、sameItemKey、displayTitle、category、brand、thumbnailUrl、bestPrice、originalPrice、priceRange、platformCount、platforms、highlights、matchLevel。"),
        ("PlatformOfferSummary", "productId、platform、price、originalPrice、shopName、productUrl、rating、sales、tags、reasons、score、title、imageUrl、brand、priceHistory、matchedPreferences、specs。"),
    ], (2400, 6960))

    add_heading(doc, "2.4 数据库和持久化", 2)
    add_table(doc, ("表", "字段", "考点"), [
        ("users", "id、username、password_hash、display_name、role、created_at", "BCrypt 密码哈希、username 唯一约束、角色默认 USER。"),
        ("chat_sessions", "session_id、user_id、title、created_at、updated_at", "会话归属、ON DELETE SET NULL、按 user_id/updated_at 建索引。"),
        ("chat_messages", "message_id、session_id、role、text、image_ids、selected_option_ids、agent_reply、created_at", "JSONB 存储 imageIds、selectedOptionIds、agentReply，支持历史卡片恢复。"),
        ("favorites", "id、user_id、product_id、title、platform、price、shop_name、brand、image_url、product_url、snapshot、created_at", "UNIQUE(user_id, product_id) 防重复收藏。"),
        ("price_alerts", "id、user_id、product_id、title、platform、target_price、triggered、last_observed_price、note、created_at", "目标价检测、triggered 状态和最近观察价格。"),
        ("user_preferences", "user_id、payload、updated_at", "JSONB 存画像，当前前端主要用本地 shared_preferences。"),
    ], (1600, 4300, 3460))
    add_bullets(doc, [
        "仓储切换用 @ConditionalOnProperty(name = \"app.persistence.store\", havingValue = \"memory\" / \"postgres\")。",
        "默认 application.yml 排除了 DataSourceAutoConfiguration 和 FlywayAutoConfiguration；postgres profile 才恢复 datasource 和 flyway。",
        "内存仓储使用 ConcurrentHashMap，适合演示和测试；Postgres 仓储适合持久化环境。"
    ])

    add_heading(doc, "2.5 后端高频问答", 2)
    add_qa(doc, "为什么 Spring Security 配置 anyRequest().permitAll，还能保护收藏和价格提醒？",
           "SecurityConfig 只关闭过滤链级别拦截，使 AI 演示链路不被登录阻断。收藏和价格提醒在 Controller 内调用 CurrentUser.require，如果 AUTH_ENABLED=true 且无有效 Bearer token，会返回 401。",
           "SecurityConfig、CurrentUser、FavoriteController、PriceAlertController。",
           "追问点是生产环境如何改：可以把 /api/auth/**、/api/health、演示公开接口放行，把收藏/价格提醒改成 filter chain 保护。")
    add_qa(doc, "JWT 怎么签发和校验？",
           "JwtService 使用 JJWT，subject 存 user.id，claims 存 username、displayName、role，issuedAt 和 expiration 使用 ttlMinutes；parse 时 verifyWith(signingKey)，异常返回 Optional.empty。",
           "JwtService.issue、JwtService.parse、AuthService.toResult。",
           "密钥不足 32 字节时源码会 padding 到 32 字节；面试要说明这是开发兜底，生产必须使用足够强的密钥。")
    add_qa(doc, "为什么图片上传要保存文件又保存 ImageStore？",
           "文件内容落盘到 uploadDir，ImageStore 只保存 imageId、storedFileName、contentType、size 等元数据。识别接口通过 imageId 查元数据，再读取文件 bytes 构造 ImagePayload。",
           "ImageController、ImageStore、RecognitionController、ChatController。")
    add_qa(doc, "如何解释 memory/postgres 双仓储？",
           "Repository 接口隔离业务，memory 版本保证本地演示和测试，postgres 版本保证持久化；Spring 条件装配让上层 Controller 和 Service 不关心具体实现。",
           "UserRepository、ChatHistoryRepository、FavoriteRepository、PriceAlertRepository 及各自实现。")

    add_heading(doc, "3. Agent、AI 和推荐检索", 1)
    add_heading(doc, "3.1 Agent 主链路", 2)
    add_numbered(doc, [
        "用户在 Flutter 聊天首页输入 text、上传 imageIds 或点击 selectedOptionIds。",
        "ChatController 校验会话和输入；图片链路先读取本地文件 bytes 调 AiRecognitionProvider。",
        "MockAgent.process / processWithRecognition 读取会话历史，合并文本、识别元数据和选项偏好。",
        "ShoppingIntentParser 解析 keyword、maxPrice、color、brand、platforms、sortBy、minRating 和偏好布尔值。",
        "CategoryResolver 把细分词归一到标准品类。",
        "CompositeProductSourceProvider 读取 public-product-offers.json，并按模式生成平台报价。",
        "RecommendationScorer 打分，RecommendationExplainer 生成 decisionSignals、evidence、risks、productAnalyses。",
        "MockAgent 输出 product_group_list 和 clarification；assistant 消息保存完整 agentReply 供历史恢复。",
    ])

    add_heading(doc, "3.2 三段式 AI Provider", 2)
    add_table(doc, ("链路", "接口/类", "默认与回退"), [
        ("图片识别", "AiRecognitionProvider -> MockRecognitionProvider / ArkRecognitionProvider / FallbackRecognitionProvider", "AI_PROVIDER=mock 用 Mock；AI_PROVIDER=ark 用 FallbackRecognitionProvider 包 Ark，失败回退 Mock。"),
        ("购物意图解析", "ShoppingIntentParser -> RuleBasedShoppingIntentParser / ArkShoppingIntentParser / FallbackShoppingIntentParser", "默认规则；Ark 失败或未配置回退规则。"),
        ("推荐解释改写", "RecommendationExplainer / ArkRecommendationExplainer", "规则先生成结构化解释；Ark 只能改写面向用户文本，失败回退规则解释。"),
        ("语音转写", "VoiceTranscriber -> MockVoiceTranscriber / ArkVoiceTranscriber / FallbackVoiceTranscriber", "VOICE_PROVIDER=mock 默认；ark 时失败回退 Mock。"),
    ], (1900, 4100, 3360))
    add_callout(doc, "AI 边界", "Ark 不允许改写 productId、platform、price、productName、score、decisionScore。源码中 ArkRecommendationExplainer.rewriteAnalyses 始终保留规则结果的 productId、platform、title、rank、score。")

    add_heading(doc, "3.3 Prompt 与结构化输出", 2)
    add_table(doc, ("Prompt", "固定字段", "防幻觉策略"), [
        ("ArkRecognitionProvider", "category、brand、model、keywords、attributes、confidence、explanation", "category 只能优先输出运动鞋、耳机、吹风机、背包、智能手表；细分词进 attributes.subCategory；confidence clamp 到 0-1。"),
        ("ArkShoppingIntentParser", "keyword、maxPrice、color、officialStore、fastDelivery、lowestPrice、highRating、highSales、brand、platforms、sortBy、minRating、needsClarification、clarificationQuestion", "platforms 白名单为京东-mock、拼多多-mock、淘宝-mock、天猫-mock；sortBy 白名单为 recommended、price_asc、price_desc、sales_desc、rating_desc；minRating 越界丢弃。"),
        ("ArkRecommendationExplainer", "summaryReason、signals、evidence、risks、analyses", "只改写 explanation、evidence、risks、strengths、weaknesses；商品身份和数值分数保留规则结果。"),
        ("ArkClient", "chat/completions 响应 content 中抽 JSON", "normalizeBaseUrl 去尾斜杠和 /chat/completions；temperature=0.1；extractJsonObject 剥离 Markdown 代码块。"),
    ], (1900, 3500, 3960))

    add_heading(doc, "3.4 多轮上下文合并", 2)
    add_table(doc, ("字段", "规则", "例子"), [
        ("品类", "当前文本明确品类 > 历史文本明确品类 > 历史识别元数据 category > 默认运动鞋。", "推荐耳机 -> 只看300以内的黑色款，仍然按耳机。"),
        ("maxPrice/color/brand/platforms/sortBy/minRating", "最近一次明确值生效，当前文本可覆盖历史值。", "500以内 -> 300以内 -> 黑色款，预算按 300。"),
        ("officialStore/fastDelivery/lowestPrice/highRating/highSales", "布尔偏好跨轮累积。", "只看官方 + 配送快，后续筛选继续保留。"),
        ("selectedOptionIds", "点击选项会转成偏好；lowest_price 未指定 sortBy 时自动价格升序。", "查看同款低价触发 price_asc。"),
        ("profile", "前端画像字段进入 ProductSearchQuery.profile，供重排/个性化扩展。", "preferredPlatforms、inferredBrands、inferredPriceMin/Max。"),
    ], (2200, 4200, 2960))

    add_heading(doc, "3.5 检索、RAG 和推荐排序", 2)
    add_table(doc, ("能力", "源码事实", "八股回答角度"), [
        ("taxonomy 归一", "CategoryResolver 读取 data/category-taxonomy.json；CategoryEntry 字段 categoryId、name、aliases、attributes。", "这是轻量本地检索；生产可替换为 RAG TopK 召回 + 受限集合选择 categoryId。"),
        ("商品源", "CompositeProductSourceProvider 支持 public-dataset-only 和 public-dataset-platforms。", "接口隔离数据源，新增真实电商 Provider 不影响 Agent 主流程。"),
        ("同款分组", "ProductOffer.sameItemKey 使用公开样例商品原始 productId，把不同平台报价聚合成 ProductGroup。", "比单纯商品列表更适合比价场景。"),
        ("规则评分", "RecommendationScorer 使用价格、评分、销量、品牌、预算、平台质量、渠道可信 7 维，归一到 0-10。", "候选集内动态阈值，使用 median、percentile、IQR，避免硬编码固定价格线。"),
        ("解释分", "RecommendationExplainer 使用 match 25%、price 25%、reputation 20%、channel 15%、risk 15%，输出 0-100 decisionScore。", "商品排序分和解释决策分不是同一套分。"),
        ("向量索引", "ProductVectorIndex 使用字典词、字符 bigram、unigram、brand/category boost token 的 TF-IDF cosine。", "无外部分词依赖，适合中文商品短文本召回。"),
        ("混合召回", "HybridRetriever 融合 vector 0.4、BM25 0.3、LLM 0.3；Ark 不可用时 0.7 vector + 0.3 BM25。", "这是独立增强检索能力，不是当前聊天主链路的默认调用。"),
        ("MMR 重排", "ResultReRanker 使用基础质量、文本/profile 匹配和 MMR 多样性；MMR_LAMBDA=0.75。", "平衡相关性与品牌/平台冗余。"),
    ], (1700, 4100, 3560))

    add_heading(doc, "3.6 Agent 高频问答", 2)
    add_qa(doc, "为什么要有 Mock Provider？",
           "Mock Provider 保证没有 Ark Key、网络异常或模型输出异常时，演示和测试仍能完整跑通。项目把 Ark 定位为增强而不是唯一依赖。",
           "MockRecognitionProvider、RuleBasedShoppingIntentParser、MockVoiceTranscriber 以及 Fallback* 类。",
           "要主动说清楚 Mock 数据不是生产能力，而是工程鲁棒性和可测性的手段。")
    add_qa(doc, "如何防止大模型幻觉影响交易结果？",
           "商品 ID、平台、价格、排序、score、decisionScore 等关键字段都来自规则和数据链路；Prompt 限定 JSON Schema，代码做白名单、范围校验和字段缺失回退。",
           "ArkShoppingIntentParser、ArkRecommendationExplainer、ArkClient.extractJsonObject。")
    add_qa(doc, "RAG 在项目里体现在哪里？",
           "当前主链路是本地 taxonomy 轻量检索归一，不是完整向量库 RAG。项目预留 ProductVectorIndex、HybridRetriever、ArkQueryDecomposer 等检索增强类，可扩展为 TopK 召回 + LLM 受限选择。",
           "CategoryResolver、category-taxonomy.json、ProductVectorIndex、HybridRetriever。",
           "不要把当前交付说成真实线上 RAG 系统；准确表述是具备 RAG 扩展路径。")
    add_qa(doc, "推荐分数怎么解释？",
           "需要区分 ProductOffer.score 和 RecommendationExplanation.decisionScore。前者由 RecommendationScorer 的 7 维商品评分归一到 0-10；后者由 RecommendationExplainer 的 5 个 decisionSignals 加权到 0-100。",
           "RecommendationScorer、RecommendationExplainer、ProductAnalysis。")

    add_heading(doc, "4. 数据、Mock 和业务边界", 1)
    add_heading(doc, "4.1 公开样例商品数据", 2)
    add_table(doc, ("项目", "内容"), [
        ("默认文件", "backend/src/main/resources/data/public-product-offers.json。"),
        ("基础商品数", "243 个。"),
        ("品类分布", "运动鞋 59、耳机 75、吹风机 17、背包 92。"),
        ("运行时平台", "京东-mock、淘宝-mock、天猫-mock、拼多多-mock。"),
        ("基础字段", "productId、category、title、platform、price、originalPrice、shopName、imageUrl、productUrl、rating、sales、brand、tags、sourceCategory、rawRating。"),
        ("生成字段", "platform、price、originalPrice、shopName、rating、sales、tags、priceHistory、sameItemKey、matchedPreferences。"),
        ("数据限制", "sales 全部为 0；35 条 rating 非 0；73 条原始 brand 为空；许可证标记 unknown。"),
    ], (2300, 7060))
    add_callout(doc, "面试边界", "平台报价、店铺名、价格历史均为代码生成的演示数据，不代表真实平台价格、库存、评价或配送服务；正式发布前需要确认公开数据授权。")

    add_heading(doc, "4.2 数据构建脚本考点", 2)
    add_bullets(doc, [
        "scripts/build_public_product_offers.py 从 jason1966/PromptCloudHQ_flipkart-products 的 flipkart_com-ecommerce_sample.csv 构建公开样例商品。",
        "文档记录原始 CSV SHA-256：56f8f699c9e847356666c2eab3c3ab1244340f6a98ad08e39ea2199ebe993ad1。",
        "筛选规则：运动鞋匹配 sports shoes / running shoes；耳机匹配 headphone / headset / earphone；吹风机匹配 hair dryer；背包匹配 backpack。",
        "确定性构建的面试价值：同输入得到相同输出，可复现、便于测试，不依赖真实平台实时波动。"
    ])

    add_heading(doc, "5. Flutter Android 客户端考点", 1)
    add_heading(doc, "5.1 路由和页面", 2)
    add_table(doc, ("路由", "页面", "职责"), [
        ("/home", "ChatScreen", "聊天首页、图片上传、语音按钮、动态建议卡、商品分组和筛选编辑。"),
        ("/login", "LoginScreen", "注册/登录入口。"),
        ("/me", "ProfileScreen", "我的页面、健康状态和电商状态入口。"),
        ("/favorites", "FavoritesScreen", "收藏商品列表。"),
        ("/price-alerts", "PriceAlertsScreen", "价格提醒列表和检测。"),
        ("/preferences", "PreferencesScreen", "用户偏好、推断画像、隐私与清空记忆。"),
    ], (1700, 2400, 5260))
    add_heading(doc, "5.1.1 Android 配置追问", 3)
    add_bullets(doc, [
        "AndroidManifest 声明了 INTERNET、CAMERA 权限，并设置 android:usesCleartextTraffic=\"true\"，便于本地 HTTP 后端联调。",
        "MainActivity 使用 FlutterActivity；release 当前使用 debug signingConfig。",
        "本地文件未声明 RECORD_AUDIO、READ_MEDIA_IMAGES、READ_EXTERNAL_STORAGE；因此不能把当前语音入口说成真实麦克风录音采集。"
    ])

    add_heading(doc, "5.2 状态管理和网络", 2)
    add_table(doc, ("点", "项目实现", "面试考点"), [
        ("API Base URL", "apiBaseUrlProvider 读取 String.fromEnvironment('EC26B_API_BASE_URL')，并去掉尾部 /。", "Flutter dart-define、环境隔离、模拟器 10.0.2.2 与 USB adb reverse。"),
        ("聊天状态", "chatControllerProvider = ChangeNotifierProvider<ChatController>。", "ChangeNotifier 的消息列表、sending 状态、会话切换、错误态。"),
        ("HTTP", "ChatApi、AuthApi、RecognitionApi、FavoriteApi、PriceAlertApi、VoiceApi、EcommerceApi。", "JSON 编解码、状态码和 code 双重判断、multipart 文件上传。"),
        ("图片", "ImagePicker.pickImage(source, imageQuality=85, maxWidth=1920, maxHeight=1920)。", "移动端图片压缩、预览、上传失败仍可发文字。"),
        ("详情页", "ProductGroupDetailScreen 接收 ProductGroup。", "页面间传参、商品详情、价格提醒、收藏、平台跳转行为事件。"),
        ("自定义绘制", "ProductThumbPainter 和价格走势 sparkline。", "CustomPainter、占位图、网络图失败 fallback。"),
    ], (1900, 3900, 3560))

    add_heading(doc, "5.3 Dart 模型字段", 2)
    add_table(doc, ("模型", "字段"), [
        ("AgentReply", "replyId、replyType、text、cards。"),
        ("ReplyCard", "cardType、title、productName、platform、price、reason、options、imageId、category、brand、model、keywords、attributes、confidence、aiProvider、fallbackUsed、explanation、recognitionId、products、platformStats、decisionScore、decisionSignals、evidence、risks、productAnalyses、intentProvider、intentFallbackUsed、explanationProvider、explanationFallbackUsed、notices、filterSummary、groups、emptyReason。"),
        ("ProductGroup", "groupId、sameItemKey、displayTitle、category、brand、thumbnailUrl、bestPrice、originalPrice、priceRange、platformCount、platforms、highlights、matchLevel。"),
        ("PlatformOfferSummary", "productId、platform、price、originalPrice、shopName、productUrl、rating、sales、tags、reasons、score、title、imageUrl、brand、priceHistory、matchedPreferences、specs。"),
        ("UserProfile", "preferredPlatforms、preferredCategories、categoryMaxBudget、decisionFactors、dislikes、inferredCategories、inferredPriceMin、inferredPriceMax、inferredPlatforms、inferredBrands、recentInterests、personalizationEnabled。"),
        ("HealthStatus", "status、app、stage、aiProvider、persistenceStore、authEnabled、ecommerceProvider、voiceProvider、timestamp。"),
    ], (2400, 6960))

    add_heading(doc, "5.4 本地记忆和个性化", 2)
    add_table(doc, ("能力", "源码事实", "面试回答"), [
        ("存储", "MemoryStore 用 shared_preferences 保存 memory_profile、memory_events、memory_onboarding_done、memory_personalization_enabled、memory_privacy_accepted。", "本地优先，用户可关闭个性化和清空记忆。"),
        ("事件上限", "saveEvents 保留最多 500 条事件。", "控制本地存储体积，避免无限增长。"),
        ("行为事件", "BehaviorEventType：search、productView、productClick、favorite、unfavorite、platformJump、priceAlertCreate、filterApply、preferenceUpdate。", "统一事件入口，避免埋点散落。"),
        ("推断画像", "ProfileEngine 使用 7 天窗口和 _decay；权重 search=1、productView=2、productClick=3、priceAlertCreate=4、favorite=5、platformJump=5、unfavorite=-2。", "近期行为权重更高，显式偏好和行为推断分层。"),
        ("请求画像", "ChatScreen._profileForRequest 去掉 personalizationEnabled 后，把 profile 随聊天请求发送。", "最小必要字段传输；关闭个性化时返回 null。"),
    ], (1800, 4300, 3260))

    add_heading(doc, "5.5 Flutter 高频问答", 2)
    add_qa(doc, "为什么用 Riverpod？",
           "项目用 Provider 注入 API 客户端，用 ChangeNotifierProvider 管聊天状态，用 StateNotifierProvider 管用户画像。测试可以 override API Provider，便于 fake 网络层。",
           "apiBaseUrlProvider、chatControllerProvider、userProfileProvider。")
    add_qa(doc, "客户端如何恢复历史卡片？",
           "后端历史消息的 assistant 记录包含完整 agentReply；ChatController.switchToSession 调 getMessages 后用 AgentReply.fromJson 恢复 ReplyCard、ProductGroup、PlatformOfferSummary 等模型。",
           "ChatController._fromHistoryMessage、ChatApi.getMessages、chat_models.dart。")
    add_qa(doc, "图片上传和识别为什么拆成两步？",
           "客户端先通过 /api/images/upload 得到 imageId，再把 imageId 放进聊天消息或识别请求。这样聊天链路只传图片引用，后端统一读取文件 bytes。",
           "ChatApi.uploadImage、ChatApi.sendMessage、RecognitionApi.recognizeImage。")
    add_qa(doc, "RecognitionApi.recognizeImage 在客户端主链路中怎么用？",
           "源码中 RecognitionApi.recognizeImage 有定义，但生产 lib 代码没有调用它；当前主图片链路是上传图片后通过聊天消息发送 imageIds，由 ChatController 在后端识别并返回卡片。",
           "RecognitionApi、ChatScreen._sendMessage、ChatController.sendMessage。",
           "面试时不要说 Flutter 主链路会先调 /api/recognition 再调聊天接口。")
    add_qa(doc, "个性化推荐如何保护隐私？",
           "原始行为事件存在手机本地；只有开启个性化时，推断出的 profile 随请求发送；用户可关闭 personalizationEnabled 或清空 MemoryStore。",
           "MemoryStore、PreferencesScreen、ChatScreen._profileForRequest。")
    add_qa(doc, "客户端网络层还有哪些不足？",
           "当前使用 package:http 直接请求，源码未设置 timeout、retry、cancel、interceptor；HealthApi.fetch 直接解析原始 JSON，不按 code/data/message 包装解析，和多数 API 客户端不一致。",
           "AuthApi、ChatApi、HealthApi。")

    add_heading(doc, "6. 测试与验收考点", 1)
    add_heading(doc, "6.1 后端测试", 2)
    add_table(doc, ("测试类", "覆盖重点"), [
        ("ChatControllerTest", "会话、消息、购物文本、预算解析、多平台商品、同款分组、超预算、多轮筛选、历史恢复、推荐解释、动态建议。"),
        ("PreferenceExtensionTests", "中英文品牌、平台、排序、最低评分、组合过滤、ProductOffer 扩展字段、PlatformStats.averagePrice。"),
        ("ArkIntegrationTests", "规则解析、Ark 未配置回退、Ark 解释回退、Ark 不改商品 ID、content type normalize。"),
        ("CategoryResolverTest", "细分品类归一、属性 schema、支持品类。"),
        ("CompositeProductSourceProviderTest / PublicDatasetProductSourceProviderTest", "商品源模式、平台变体、筛选和排序。"),
        ("AuthControllerTest", "注册、登录、当前用户接口。"),
        ("ImageControllerTest / RecognitionControllerTest", "上传、识别、识别修正、错误处理。"),
        ("FavoriteControllerTest / PriceAlertControllerTest / VoiceControllerTest", "收藏、价格提醒、语音转写。"),
    ], (2600, 6760))
    add_bullets(doc, [
        "docs/07 记录后端 mvn test：142 tests，0 failures，0 errors。",
        "测试框架为 JUnit 5、Spring Boot Test、MockMvc。",
        "面试重点不是背数量，而是说明覆盖了主链路、异常链路、回退链路和历史恢复。"
    ])

    add_heading(doc, "6.2 Flutter 测试", 2)
    add_table(doc, ("类别", "覆盖重点"), [
        ("AppBar/Profile", "历史入口、我的入口、图片/语音/发送按钮、ProfileScreen 分区。"),
        ("聊天", "发送文字、追问卡、商品分组卡、识别卡、识别修正、历史恢复、语音填充。"),
        ("商品卡扩展", "品牌徽章、价格走势、偏好命中徽章、平台报价、平均价、filterSummary、动态建议选项。"),
        ("详情页", "购买判断、评价概览、平台比价、价格提醒、收藏、返回聊天。"),
        ("模型解析", "ReplyCard、PlatformOfferSummary、EcommerceStatus 等 JSON 解析。"),
    ], (2300, 7060))
    add_bullets(doc, [
        "docs/07 记录 Flutter analyze：0 error / 0 warning，33 条 info。",
        "docs/07 记录 Flutter test：41 tests，全部通过；其中 37 条 widget 测试，4 条模型解析测试。",
        "测试使用 Riverpod override 和 Fake ChatApi / Fake RecognitionApi / Fake VoiceApi。"
    ])

    add_heading(doc, "7. 现有八股资料如何迁移到项目", 1)
    add_table(doc, ("八股分类", "已有覆盖", "项目对应讲法"), [
        ("Java 后端语言基础", "44 个问题", "record、List/Map、泛型、注解、反射、final/static、接口和实现可结合 DTO、Repository、Provider。"),
        ("并发编程", "1 个问题", "ConcurrentHashMap 在 ImageStore、RecognitionStore、内存仓储中使用；要补线程安全集合、原子性、并发容器。"),
        ("JVM", "2 个问题", "Java 21、Spring Boot 启动、对象生命周期和 GC 可结合服务端常驻进程。"),
        ("MySQL", "76/77 个文件", "项目用 PostgreSQL，但索引、事务、MVCC、JSONB、唯一约束、外键、分页思想可迁移；回答时不要把 PostgreSQL 说成 MySQL。"),
        ("Redis", "37 个问题", "项目当前未使用 Redis；可作为后续缓存会话、价格提醒状态、限流、分布式锁的扩展方案。"),
        ("Spring 体系", "28 个问题", "Spring Boot 自动配置、@RestController、@Bean、@Primary、@Value、@ConditionalOnProperty、MockMvc。"),
        ("计算机网络", "3 个问题", "REST/JSON、Multipart、Authorization Bearer、移动端 baseUrl、HTTP 状态码。"),
        ("操作系统与 Linux", "1 个问题", "上传文件落盘、路径 normalize、部署脚本、日志和端口排查。"),
        ("分布式/高并发", "1 个问题", "当前单体演示；可讲后续用 Redis/消息队列/定时任务扩展价格提醒和真实平台抓取。"),
        ("服务治理与容器", "5 个问题", "当前无容器化交付；可讲健康检查 /api/health、配置外置化和未来 Docker 部署。"),
        ("安全与权限", "空分类", "本项目有 JWT、BCrypt、AUTH_ENABLED、Controller 内鉴权、密钥不入库等内容，建议补成重点。"),
        ("设计模式与代码质量", "空分类", "策略模式、装饰/回退、仓储模式、DTO/record、接口隔离、测试替身。"),
        ("工程工具", "空分类", "Maven、Flutter CLI、Flyway、dart analyze、自动化测试、确定性数据脚本。"),
    ], (1900, 1700, 5760))

    add_heading(doc, "8. 按主题背诵的项目八股", 1)
    qa_items = [
        ("Spring Boot 自动配置为什么要排除 DataSourceAutoConfiguration？",
         "默认 memory 模式不需要数据源，排除 DataSourceAutoConfiguration 和 FlywayAutoConfiguration 可以让后端无数据库也能启动；postgres profile 再恢复 datasource 和 flyway。",
         "application.yml 的 spring.autoconfigure.exclude 和 postgres profile。"),
        ("@Primary 在项目中解决什么问题？",
         "当存在多个同类型 Bean 时，@Primary 指定默认注入对象。本项目 CompositeProductSourceProvider 是 ProductSourceProvider 的主实现，AiConfig/VoiceConfig 也用 @Primary 装配当前 Provider。",
         "CompositeProductSourceProvider、AiConfig.aiRecognitionProvider、AiConfig.shoppingIntentParser、VoiceConfig.voiceTranscriber。"),
        ("为什么 Controller 里还要做参数校验？",
         "项目没有把所有请求体都用 Bean Validation 注解覆盖，Controller 对空文件、空标题、空消息、无效 targetPrice 等做显式校验，返回统一 ApiResponse 错误。",
         "ImageController、ChatController、PriceAlertController、FavoriteController。"),
        ("如何解释 JSONB 的选择？",
         "chat_messages 的 image_ids、selected_option_ids、agent_reply 以及 user_preferences.payload 用 JSONB，适合保存结构变化快的卡片和画像数据，减少频繁改表。",
         "V1__init.sql。"),
        ("为什么 product_group_list 比 product_list 更适合比价？",
         "同款分组把不同平台报价聚合到同一 ProductGroup，用户看到 bestPrice、priceRange、platformCount 和平台报价明细，而不是多条重复商品。",
         "MockAgent.groupProducts、ProductGroup、PlatformOfferSummary。"),
        ("价格提醒是实时的吗？",
         "当前不是实时真实平台提醒。POST /api/price-alerts/check 遍历用户 alert，在 Mock 商品源里按 productId 找当前样例价格并更新 triggered。",
         "PriceAlertController.check。"),
        ("语音能力做到什么程度？",
         "后端有 /api/voice/transcribe，默认 MockVoiceTranscriber；VOICE_PROVIDER=ark 时走 FallbackVoiceTranscriber 包 ArkVoiceTranscriber。Flutter 当前发送演示音频字节，不做真实录音采集。",
         "VoiceController、VoiceConfig、VoiceApi。"),
        ("为什么同一个项目既有规则解析又有 Ark 解析？",
         "规则解析可测试、稳定、可控；Ark 解析提升表达覆盖和灵活性。FallbackShoppingIntentParser 让 Ark 增强失败时不影响主链路。",
         "RuleBasedShoppingIntentParser、ArkShoppingIntentParser、FallbackShoppingIntentParser。"),
        ("如何处理 Ark 返回非 JSON？",
         "ArkClient 从 choices/0/message/content 取文本，extractJsonObject 剥离 ```json 代码块并截取最外层大括号；解析失败抛异常，由 fallback 捕获。",
         "ArkClient.chatJson、ArkClient.extractJsonObject。"),
        ("为什么要做 content-type normalize？",
         "上传或系统可能把图片声明成 application/octet-stream，ArkRecognitionProvider 通过 JPEG/PNG/WebP 头字节判断真实 MIME，提高视觉模型接受成功率。",
         "ArkRecognitionProvider.normalizeContentType。"),
        ("多轮筛选如何避免丢失上下文？",
         "MockAgent.mergeContext 扫描会话历史，按字段合并。品类优先级、数值字段覆盖、偏好布尔累积三条规则保证短句筛选能继承上一轮意图。",
         "MockAgent.mergeContext、docs/03。"),
        ("如何讲项目中的可解释推荐？",
         "先说商品评分由 RecommendationScorer 计算，再说 RecommendationExplainer 输出五个 decisionSignals、evidence、risks、productAnalyses，前端用卡片展示。",
         "RecommendationScorer、RecommendationExplainer、chat_models.dart。"),
        ("为什么没有真实电商 API？",
         "真实电商接口涉及登录、风控、库存、价格波动和合规问题。当前赛题/项目阶段用公开数据集和 Mock 平台验证端到端能力，文档明确不宣称真实平台数据。",
         "docs/05、README 当前边界。"),
        ("客户端如何处理后端错误？",
         "API 层判断 HTTP 200 且 body['code']==0；失败抛异常。ChatController 捕获发送失败后移除 loading 消息并追加错误文案。",
         "ChatApi、AuthApi、ChatController.sendTextMessage。"),
        ("Riverpod override 在测试中有什么价值？",
         "可以用 Fake ChatApi / RecognitionApi / VoiceApi 替换真实网络，widget 测试只验证 UI 状态和模型解析，不依赖后端服务。",
         "app/test/app_test.dart、chatControllerProvider。"),
        ("本地记忆如何影响推荐？",
         "ChatScreen 发送消息时调用 _profileForRequest，把开启个性化后的 UserProfile JSON 去掉 personalizationEnabled 后传给后端 profile 字段。",
         "ChatScreen._profileForRequest、ChatApi.sendMessage、ProductSearchQuery.profile。"),
        ("怎么回答项目不足？",
         "真实平台价格/库存/评价、真实语音识别、图片/识别结果持久化、完整线上账号安全、部署容器化、数据授权确认不在当前交付范围；这些都有明确后续计划。",
         "README 当前边界、docs/08。"),
    ]
    for q, a, p in qa_items:
        add_qa(doc, q, a, p)

    add_heading(doc, "9. 答辩速查", 1)
    add_heading(doc, "9.1 5 分钟演示顺序", 2)
    add_numbered(doc, [
        "拍照或选择商品图片，展示上传预览和识别元数据。",
        "返回 product_group_list 和 category 差异化动态建议卡。",
        "点击“查看同款低价”或输入“只看300以内的黑色款”。",
        "展示同款分组、最低价、价格区间、平台报价、命中偏好和价格走势。",
        "进入商品详情，展示购买判断、评价概览、平台比价。",
        "演示收藏和价格提醒，说明当前是 Mock 商品源。"
    ])

    add_heading(doc, "9.2 10 分钟设计讲解提纲", 2)
    add_table(doc, ("时间", "讲什么", "必须出现的关键词"), [
        ("1 分钟", "项目定位", "聊天式购物 Agent、Flutter Android、Spring Boot Java 21、Mock 闭环 + Ark 增强。"),
        ("2 分钟", "整体架构", "api/chat/ai/product/image/auth、REST/JSON/Multipart、memory/postgres。"),
        ("3 分钟", "AI Pipeline", "识别/意图/解释三条链路、JSON Schema、白名单、fallback、notices。"),
        ("2 分钟", "检索推荐", "CategoryResolver、public-dataset-platforms、sameItemKey、7 维评分、5 个 decisionSignals。"),
        ("1 分钟", "客户端", "Riverpod、GoRouter、shared_preferences、profile、CustomPainter、历史恢复。"),
        ("1 分钟", "测试验收", "后端 142 tests、Flutter 41 tests、Ark 回退、空状态、多轮筛选。"),
    ], (1200, 3300, 4860))

    add_heading(doc, "9.3 评委追问雷区", 2)
    add_table(doc, ("不要这么说", "应该这么说"), [
        ("我们接入了真实电商比价。", "当前使用公开样例数据和代码生成的四个平台 Mock 报价，不代表真实平台数据。"),
        ("这是完整 RAG。", "当前主链路是本地 taxonomy 轻量检索归一；代码有向量/混合召回扩展能力，可演进为 RAG。"),
        ("AI 决定商品价格和排序。", "价格来自数据/平台变体生成，排序和分数由规则链路控制，AI 只做识别/解析增强和解释改写。"),
        ("Spring Security 已经保护所有接口。", "过滤链 permitAll，收藏和价格提醒由 Controller 内 CurrentUser.require 控制；生产可迁移到过滤链鉴权。"),
        ("语音已经真实录音识别。", "当前 Flutter 发送演示音频字节，后端接口和 Provider 回退已打通，真实录音采集不在当前范围。"),
        ("样例数据授权没问题。", "文档记录来源许可证标记 unknown，正式发布或商业使用前必须确认授权。"),
    ], (3300, 6060))

    add_heading(doc, "9.4 主动说明的工程风险", 2)
    add_table(doc, ("风险点", "准确表述", "改进方向"), [
        ("HybridRetriever / ArkQueryDecomposer / ResultReRanker", "这些类存在于 product 包，但未接入运行时主调用链。", "接入 ProductSourceProvider 或 MockAgent 推荐链路，并补端到端测试。"),
        ("AUTH_ENABLED=false + postgres", "CurrentUser 会返回 userId=0；Postgres favorites/price_alerts 有 user_id 外键，迁移脚本没有插入 id=0 的用户。", "postgres 环境启用真实登录，或迁移脚本创建 demo 用户。"),
        ("接口用户隔离", "SecurityConfig 放行所有请求，/api/chat/**、/api/images/upload、/api/recognition 没有用户隔离。", "把会话、图片、识别结果绑定用户，并迁移到过滤链鉴权。"),
        ("图片/识别持久化", "ImageStore、RecognitionStore 是内存存储，数据库迁移没有图片元数据和识别结果表。", "新增 image_metadata、recognition_results 表并清理孤儿文件。"),
        ("登录态恢复", "AuthController 返回 token，但 Flutter 当前没有实现 token 落盘后的重启恢复。", "用安全存储保存 token，启动时调用 /api/auth/me 恢复。"),
        ("VoiceApi contentType", "VoiceApi.transcribeBytes 有 contentType 参数，但 MultipartFile.fromBytes 调用未传 content type。", "构造 MultipartFile 时设置 MediaType，并补测试。"),
    ], (2300, 4300, 2760))

    add_heading(doc, "10. 源码定位附录", 1)
    add_table(doc, ("想查的问题", "文件"), [
        ("后端依赖版本", "backend/pom.xml。"),
        ("服务配置和 profile", "backend/src/main/resources/application.yml。"),
        ("数据库表结构", "backend/src/main/resources/db/migration/V1__init.sql。"),
        ("API 路由", "backend/src/main/java/com/ec26b/shoppingagent/api/*.java；docs/04-API设计.md。"),
        ("Agent 主流程", "backend/src/main/java/com/ec26b/shoppingagent/chat/MockAgent.java。"),
        ("AI Provider 装配", "backend/src/main/java/com/ec26b/shoppingagent/config/AiConfig.java。"),
        ("Ark 调用", "backend/src/main/java/com/ec26b/shoppingagent/ai/ArkClient.java。"),
        ("意图解析", "backend/src/main/java/com/ec26b/shoppingagent/product/RuleBasedShoppingIntentParser.java；ArkShoppingIntentParser.java。"),
        ("检索和推荐", "CategoryResolver.java、CompositeProductSourceProvider.java、PublicDatasetProductSourceProvider.java、RecommendationScorer.java、RecommendationExplainer.java。"),
        ("客户端 API", "app/lib/core/network/api_client.dart；app/lib/features/**/**/*_api.dart。"),
        ("客户端模型", "app/lib/features/chat/chat_models.dart；auth_models.dart；favorite_models.dart；price_alert_models.dart。"),
        ("客户端状态", "app/lib/features/chat/chat_controller.dart；app/lib/features/memory/user_profile.dart。"),
        ("客户端路由", "app/lib/router/app_router.dart。"),
        ("项目答辩口径", "docs/10-AI使用总结.md；docs/12-答辩材料.md。"),
        ("已有八股资料", "八股/README.md；八股/questions.js；八股各分类目录。"),
    ], (2500, 6860))

    add_heading(doc, "11. 最后一页背诵清单", 1)
    add_bullets(doc, [
        "一句话：识价镜是聊天式购物 Agent，不是普通商品列表应用。",
        "四个核心创新：三段式 AI 编排、taxonomy 检索归一、多轮 7 维自然语言筛选、动态建议卡。",
        "三个边界：公开样例数据不是真实电商；当前主链路不是完整 RAG；语音是接口和 Mock/Ark Provider 链路，不是真实录音采集。",
        "两个评分：ProductOffer.score 是 0-10 商品排序分；decisionScore 是 0-100 推荐解释分。",
        "一个安全点：Spring Security permitAll 是演示友好选择，收藏和价格提醒由 CurrentUser.require 做 Controller 内鉴权。",
        "一个测试点：后端和 Flutter 都有覆盖主链路、异常链路、回退链路和 UI 恢复的自动化测试记录。",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
